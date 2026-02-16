"""One-time script to export brand voice guide from .docx to plain text."""

import sys
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from config.settings import BRAND_VOICE_FILE


def export_docx_to_txt(docx_path: str, output_path: Path = BRAND_VOICE_FILE):
    """Convert a .docx brand voice guide to plain text."""
    doc = Document(docx_path)
    lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
        else:
            lines.append("")

    # Also extract table content (the brand voice guide has tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")

    content = "\n".join(lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding="utf-8")
    print(f"Brand voice guide exported to: {output_path}")
    print(f"Character count: {len(content)}")
    print(f"Word count: {len(content.split())}")


if __name__ == "__main__":
    default_source = (
        "/Users/madhur/Documents/Joyce Happy Journey/"
        "Happy_Journey_with_Joyce_Brand_Voice_Guide.docx"
    )
    source = sys.argv[1] if len(sys.argv) > 1 else default_source
    export_docx_to_txt(source)
