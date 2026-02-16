"""Streamlit web chat interface for Joyce's content creation."""

import io
import shutil
import streamlit as st
from docx import Document as DocxDocument

from config.settings import BRAND_VOICE_FILE, BRAND_VOICE_BACKUP
from core.content_engine import ContentEngine
from core.image_generator import ImageGenerator
from core.image_editor import process_uploaded_image
from core.memory import save_approval, build_preference_summary
from utils.formatting import parse_dual_options, format_clean_copy, count_hashtags


def init_session_state():
    """Initialize session state variables."""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "engine" not in st.session_state:
        st.session_state.engine = ContentEngine()
    if "current_result" not in st.session_state:
        st.session_state.current_result = None
    if "conversation_history" not in st.session_state:
        st.session_state.conversation_history = []


def render_sidebar():
    """Render the sidebar with controls."""
    with st.sidebar:
        st.title("Content Studio")
        st.markdown("---")

        # Platform selector
        platform = st.radio(
            "Platform",
            ["Instagram", "Rednote", "Both"],
            index=0,
            key="platform_select",
        )
        st.session_state.platform = platform.lower()

        st.markdown("---")

        # Quick actions
        st.subheader("Quick Actions")
        if st.button("New Post", use_container_width=True):
            st.session_state.mode = "post"
        if st.button("Weekly Calendar", use_container_width=True):
            st.session_state.mode = "calendar"

        st.markdown("---")

        # Image upload
        st.subheader("Upload Photo")
        uploaded_file = st.file_uploader(
            "Upload your photo for text overlay",
            type=["jpg", "jpeg", "png"],
            key="photo_upload",
        )
        if uploaded_file:
            st.session_state.uploaded_image = uploaded_file.read()
            st.image(st.session_state.uploaded_image, caption="Uploaded photo", width=200)

        st.markdown("---")

        # Brand voice upload
        st.subheader("Brand Voice")
        voice_file = st.file_uploader(
            "Upload new brand voice (.docx)",
            type=["docx"],
            key="voice_upload",
        )
        if voice_file:
            _update_brand_voice(voice_file)

        # Show preference summary
        pref_summary = build_preference_summary()
        if pref_summary:
            st.markdown("---")
            st.subheader("Learned Preferences")
            st.caption(pref_summary)


def render_chat():
    """Render the main chat interface."""
    # Display chat history
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("image"):
                st.image(msg["image"], width=400)

    # Chat input
    user_input = st.chat_input("What would you like to create?")
    if user_input:
        _handle_user_input(user_input)


def _handle_user_input(user_input: str):
    """Process user input and generate content."""
    # Add user message
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)

    # Check if this is a calendar request
    is_calendar = any(
        word in user_input.lower()
        for word in ["calendar", "week", "weekly", "schedule", "plan"]
    )

    platform = st.session_state.get("platform", "instagram")

    with st.chat_message("assistant"):
        if is_calendar:
            _handle_calendar_request(user_input, platform)
        else:
            _handle_post_request(user_input, platform)


def _handle_post_request(topic: str, platform: str):
    """Generate a single post with 2 options."""
    engine = st.session_state.engine

    with st.spinner("Generating 2 content options..."):
        result = engine.generate_post(platform=platform, topic=topic)

    st.session_state.current_result = result

    # Display Option A
    option_a = result["option_a"]
    st.subheader("Option A (Reflective)")
    if option_a.get("content_type"):
        st.caption(f"Content type: {option_a['content_type']}")
    st.markdown(option_a.get("caption") or option_a.get("raw", ""))
    if option_a.get("hashtags"):
        st.code(option_a["hashtags"])
    if option_a.get("visual"):
        st.info(f"Visual suggestion: {option_a['visual']}")

    st.markdown("---")

    # Display Option B
    option_b = result["option_b"]
    st.subheader("Option B (Direct)")
    if option_b.get("content_type"):
        st.caption(f"Content type: {option_b['content_type']}")
    st.markdown(option_b.get("caption") or option_b.get("raw", ""))
    if option_b.get("hashtags"):
        st.code(option_b["hashtags"])
    if option_b.get("visual"):
        st.info(f"Visual suggestion: {option_b['visual']}")

    # Handle uploaded image
    uploaded_image = st.session_state.get("uploaded_image")
    if uploaded_image:
        hook_line = (option_a.get("caption", "") or "").split("\n")[0][:80]
        try:
            edited_image = process_uploaded_image(uploaded_image, overlay_text=hook_line)
            st.image(edited_image, caption="Your photo with text overlay", width=400)
            st.session_state.edited_image = edited_image
        except Exception as e:
            st.warning(f"Image editing issue: {e}")

    # Action buttons
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button("Pick A", key=f"pick_a_{len(st.session_state.messages)}"):
            _approve_option(result, "a", platform, topic)
    with col2:
        if st.button("Pick B", key=f"pick_b_{len(st.session_state.messages)}"):
            _approve_option(result, "b", platform, topic)
    with col3:
        if st.button("Regenerate", key=f"regen_{len(st.session_state.messages)}"):
            st.rerun()
    with col4:
        if st.button("Generate Image", key=f"gen_img_{len(st.session_state.messages)}"):
            _generate_image(result, platform)

    # Store assistant message
    response_text = (
        f"**Option A (Reflective):**\n{option_a.get('caption', option_a.get('raw', ''))}\n\n"
        f"**Option B (Direct):**\n{option_b.get('caption', option_b.get('raw', ''))}"
    )
    st.session_state.messages.append({"role": "assistant", "content": response_text})


def _handle_calendar_request(user_input: str, platform: str):
    """Generate a weekly content calendar."""
    engine = st.session_state.engine

    # Parse themes from input
    themes = [t.strip() for t in user_input.replace("calendar", "").replace("weekly", "").split(",") if t.strip()]
    if not themes:
        themes = [user_input]

    platforms = [platform] if platform != "both" else ["instagram", "rednote"]

    with st.spinner("Generating your weekly calendar... This may take a minute."):
        calendar_text = engine.generate_calendar(
            platforms=platforms,
            themes=themes,
            num_posts=5,
        )

    st.markdown(calendar_text)

    # Download button
    st.download_button(
        label="Download Calendar",
        data=calendar_text,
        file_name="weekly_calendar.md",
        mime="text/markdown",
    )

    st.session_state.messages.append({"role": "assistant", "content": calendar_text})


def _approve_option(result: dict, option: str, platform: str, topic: str):
    """Approve an option and save to memory."""
    option_data = result[f"option_{option}"]
    clean = format_clean_copy(option_data)

    st.success("Approved! Here's your final copy:")
    st.code(clean, language=None)

    style = "reflective" if option == "a" else "direct"
    save_approval(
        platform=platform,
        topic=topic,
        chosen_option=option.upper(),
        style_used=style,
        content_type=option_data.get("content_type", "feed_post"),
        caption=option_data.get("caption", clean),
        hashtag_count=count_hashtags(option_data.get("hashtags", "")),
    )


def _generate_image(result: dict, platform: str):
    """Generate an AI image for the post."""
    try:
        engine = st.session_state.engine
        img_gen = ImageGenerator()
        caption = result["option_a"]["caption"] or result["option_a"]["raw"]

        with st.spinner("Generating AI image..."):
            image_bytes, dalle_prompt = img_gen.generate_from_caption(
                caption, platform, engine
            )

        st.image(image_bytes, caption="AI-generated image", width=400)
        st.caption(f"Prompt used: {dalle_prompt}")
    except Exception as e:
        st.error(f"Image generation failed: {e}")


def _update_brand_voice(voice_file):
    """Update the brand voice guide from uploaded .docx."""
    try:
        docx_doc = DocxDocument(io.BytesIO(voice_file.read()))
        lines = []
        for paragraph in docx_doc.paragraphs:
            text = paragraph.text.strip()
            lines.append(text if text else "")
        for table in docx_doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
            lines.append("")

        new_content = "\n".join(lines)

        if BRAND_VOICE_FILE.exists():
            shutil.copy2(BRAND_VOICE_FILE, BRAND_VOICE_BACKUP)

        BRAND_VOICE_FILE.write_text(new_content, encoding="utf-8")
        st.success(
            f"Brand voice updated! ({len(new_content.split())} words). "
            "Next generation will use the new guide."
        )
    except Exception as e:
        st.error(f"Error processing document: {e}")


def main():
    """Main Streamlit app entry point."""
    st.set_page_config(
        page_title="Happy Journey with Joyce — Content Studio",
        page_icon="🌿",
        layout="wide",
    )

    st.title("Happy Journey with Joyce")
    st.caption("Content Studio — Powered by Sora")

    init_session_state()
    render_sidebar()
    render_chat()


if __name__ == "__main__":
    main()
