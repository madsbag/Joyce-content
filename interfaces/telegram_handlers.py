"""Telegram bot conversation handlers and command logic."""

import io
import logging
import shutil
from pathlib import Path

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ContextTypes, ConversationHandler
from docx import Document as DocxDocument

from config.settings import BRAND_VOICE_FILE, BRAND_VOICE_BACKUP
from core.content_engine import ContentEngine
from core.image_generator import ImageGenerator
from core.image_editor import process_uploaded_image
from core.memory import save_approval
from utils.formatting import parse_dual_options, format_clean_copy, count_hashtags

logger = logging.getLogger(__name__)

# Conversation states
POST_PLATFORM = 0
POST_TOPIC = 1
POST_IMAGE_CHOICE = 2
POST_PHOTO_UPLOAD = 3
POST_APPROVE = 4
POST_REVISE = 5
CAL_THEMES = 10
CAL_PLATFORM = 11
CAL_FREQUENCY = 12
CAL_NAVIGATE = 13
VOICE_UPLOAD = 20

# Shared engine instances
_content_engine = None
_image_generator = None


def _get_content_engine() -> ContentEngine:
    global _content_engine
    if _content_engine is None:
        _content_engine = ContentEngine()
    return _content_engine


def _get_image_generator() -> ImageGenerator:
    global _image_generator
    if _image_generator is None:
        _image_generator = ImageGenerator()
    return _image_generator


# ============================================================
# /start and /help
# ============================================================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Welcome message with main menu."""
    keyboard = [
        [InlineKeyboardButton("Create a Post", callback_data="create_post")],
        [InlineKeyboardButton("Weekly Calendar", callback_data="weekly_calendar")],
        [InlineKeyboardButton("Help", callback_data="show_help")],
    ]
    await update.message.reply_text(
        "Hi Joyce! I'm Sora, your content partner.\n\n"
        "What would you like to create today?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Show help information."""
    text = (
        "*Available Commands:*\n\n"
        "/post — Create a single social media post\n"
        "/calendar — Generate a weekly content calendar\n"
        "/update\\_voice — Upload a new brand voice document\n"
        "/help — Show this help message\n"
        "/cancel — Cancel current operation\n\n"
        "*Free-form mode:*\n"
        "Just type naturally! For example:\n"
        "• \"Write a post about letting go of perfectionism\"\n"
        "• \"Rednote post about BodyTalk in Chinese\"\n"
        "• \"Give me 3 post ideas about career transitions\"\n\n"
        "I'll generate 2 options for every post so you can pick your favorite."
    )
    await (update.message or update.callback_query.message).reply_text(
        text, parse_mode="Markdown"
    )


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Cancel the current conversation."""
    context.user_data.clear()
    await update.message.reply_text("Cancelled. Send /start to begin again.")
    return ConversationHandler.END


# ============================================================
# Post creation flow
# ============================================================

async def post_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start the post creation flow — ask for platform."""
    keyboard = [
        [InlineKeyboardButton("Instagram", callback_data="platform_instagram")],
        [InlineKeyboardButton("Rednote", callback_data="platform_rednote")],
        [InlineKeyboardButton("Both", callback_data="platform_both")],
    ]

    message = update.callback_query.message if update.callback_query else update.message
    if update.callback_query:
        await update.callback_query.answer()

    await message.reply_text(
        "Which platform is this for?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return POST_PLATFORM


async def post_platform(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle platform selection, ask for topic."""
    query = update.callback_query
    await query.answer()

    platform = query.data.replace("platform_", "")
    context.user_data["platform"] = platform

    platform_label = {"instagram": "Instagram", "rednote": "Rednote", "both": "Both platforms"}
    await query.message.reply_text(
        f"Platform: {platform_label.get(platform, platform)}\n\n"
        "What's the topic or theme? You can be brief or detailed.\n\n"
        "Examples:\n"
        "• \"dealing with change\"\n"
        "• \"a post about how uncertainty in midlife is normal, "
        "inspired by a client conversation about career shifts\""
    )
    return POST_TOPIC


async def post_topic(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle topic input, ask about image."""
    context.user_data["topic"] = update.message.text

    keyboard = [
        [InlineKeyboardButton("Upload my photo", callback_data="image_upload")],
        [InlineKeyboardButton("Generate an image", callback_data="image_generate")],
        [InlineKeyboardButton("Skip image", callback_data="image_skip")],
    ]
    await update.message.reply_text(
        "Want to add a visual?\n\n"
        "You can upload your own photo (I'll add text overlay and branding), "
        "or I can generate one with AI.",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return POST_IMAGE_CHOICE


async def post_image_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle image choice."""
    query = update.callback_query
    await query.answer()

    choice = query.data.replace("image_", "")
    context.user_data["image_choice"] = choice

    if choice == "upload":
        await query.message.reply_text("Send me your photo and I'll make it post-ready.")
        return POST_PHOTO_UPLOAD
    else:
        # Generate content (with or without AI image)
        await query.message.reply_text("Generating your content... This may take a moment.")
        return await _generate_and_send(update, context)


async def post_photo_upload(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle photo upload from Joyce."""
    if update.message.photo:
        photo = update.message.photo[-1]  # Highest resolution
    elif update.message.document:
        photo = update.message.document
    else:
        await update.message.reply_text("Please send a photo. Try again or /cancel.")
        return POST_PHOTO_UPLOAD

    file = await photo.get_file()
    image_bytes = await file.download_as_bytearray()
    context.user_data["uploaded_image"] = bytes(image_bytes)

    await update.message.reply_text("Photo received! Generating your content...")
    return await _generate_and_send(update, context)


async def _generate_and_send(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Generate content and send both options to Joyce."""
    engine = _get_content_engine()
    platform = context.user_data["platform"]
    topic = context.user_data["topic"]

    # Generate dual-option content
    result = engine.generate_post(platform=platform, topic=topic)
    context.user_data["result"] = result
    context.user_data["conversation_history"] = [
        {"role": "user", "content": f"Topic: {topic}"},
        {"role": "assistant", "content": result["raw"]},
    ]

    message = update.callback_query.message if update.callback_query else update.message

    # Send Option A
    option_a_text = _format_option_preview(result["option_a"], "Option A (Reflective)")
    await message.reply_text(option_a_text, parse_mode="Markdown")

    # Send Option B
    option_b_text = _format_option_preview(result["option_b"], "Option B (Direct)")
    await message.reply_text(option_b_text, parse_mode="Markdown")

    # Handle images
    image_choice = context.user_data.get("image_choice", "skip")

    if image_choice == "generate":
        try:
            img_gen = _get_image_generator()
            caption_for_image = result["option_a"]["caption"] or result["option_a"]["raw"]
            image_bytes, dalle_prompt = img_gen.generate_from_caption(
                caption_for_image, platform, engine
            )
            context.user_data["generated_image"] = image_bytes
            await message.reply_photo(
                photo=image_bytes,
                caption="AI-generated image for your post",
            )
        except Exception as e:
            logger.error(f"Image generation failed: {e}")
            await message.reply_text(
                "Image generation encountered an issue. You can still use the text content."
            )

    elif image_choice == "upload" and context.user_data.get("uploaded_image"):
        try:
            hook_line = (result["option_a"]["caption"] or "").split("\n")[0][:80]
            edited_image = process_uploaded_image(
                context.user_data["uploaded_image"],
                overlay_text=hook_line,
            )
            context.user_data["edited_image"] = edited_image
            await message.reply_photo(
                photo=edited_image,
                caption="Your photo with text overlay",
            )
        except Exception as e:
            logger.error(f"Image editing failed: {e}")
            await message.reply_text(
                "Image editing encountered an issue. You can still use the original photo."
            )

    # Approval buttons
    keyboard = [
        [
            InlineKeyboardButton("Pick A", callback_data="approve_a"),
            InlineKeyboardButton("Pick B", callback_data="approve_b"),
        ],
        [
            InlineKeyboardButton("Revise A", callback_data="approve_revise_a"),
            InlineKeyboardButton("Revise B", callback_data="approve_revise_b"),
        ],
        [
            InlineKeyboardButton("Regenerate", callback_data="approve_regenerate"),
            InlineKeyboardButton("Cancel", callback_data="approve_cancel"),
        ],
    ]
    await message.reply_text(
        "Which option do you prefer?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return POST_APPROVE


async def post_approve(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle approval, revision, or regeneration."""
    query = update.callback_query
    await query.answer()

    action = query.data.replace("approve_", "")
    result = context.user_data.get("result", {})

    if action == "a" or action == "b":
        # Approve and send clean copy
        option = result.get(f"option_{action}", {})
        clean = format_clean_copy(option)
        context.user_data["approved_option"] = action.upper()

        await query.message.reply_text(
            f"Here's your final post — ready to copy and paste:\n\n"
            f"```\n{clean}\n```",
            parse_mode="Markdown",
        )

        # Save to preference memory
        style = "reflective" if action == "a" else "direct"
        save_approval(
            platform=context.user_data["platform"],
            topic=context.user_data["topic"],
            chosen_option=action.upper(),
            style_used=style,
            content_type=option.get("content_type", "feed_post"),
            caption=option.get("caption", clean),
            hashtag_count=count_hashtags(option.get("hashtags", "")),
            revision_notes=context.user_data.get("revision_notes_list", []),
        )

        # Also send image if available
        img = context.user_data.get("edited_image") or context.user_data.get("generated_image")
        if img:
            await query.message.reply_photo(photo=img, caption="Your post image")

        context.user_data.clear()
        return ConversationHandler.END

    elif action.startswith("revise_"):
        option_letter = action.replace("revise_", "")
        context.user_data["revising_option"] = option_letter
        await query.message.reply_text(
            f"What should I change about Option {option_letter.upper()}?\n\n"
            "Examples: \"make it shorter\", \"more personal\", "
            "\"add a question at the end\", \"less formal\""
        )
        return POST_REVISE

    elif action == "regenerate":
        await query.message.reply_text("Regenerating from scratch...")
        return await _generate_and_send(update, context)

    elif action == "cancel":
        context.user_data.clear()
        await query.message.reply_text("Cancelled. Send /start to begin again.")
        return ConversationHandler.END


async def post_revise_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle revision feedback and regenerate."""
    revision_note = update.message.text
    revision_list = context.user_data.get("revision_notes_list", [])
    revision_list.append(revision_note)
    context.user_data["revision_notes_list"] = revision_list

    engine = _get_content_engine()
    platform = context.user_data["platform"]
    topic = context.user_data["topic"]

    await update.message.reply_text("Revising... This may take a moment.")

    result = engine.generate_post(
        platform=platform,
        topic=topic,
        revision_notes=revision_note,
        conversation_history=context.user_data.get("conversation_history"),
    )
    context.user_data["result"] = result

    # Update conversation history
    history = context.user_data.get("conversation_history", [])
    history.append({"role": "user", "content": f"Revision: {revision_note}"})
    history.append({"role": "assistant", "content": result["raw"]})
    context.user_data["conversation_history"] = history

    # Re-send options
    option_a_text = _format_option_preview(result["option_a"], "Option A (Reflective) — Revised")
    await update.message.reply_text(option_a_text, parse_mode="Markdown")

    option_b_text = _format_option_preview(result["option_b"], "Option B (Direct) — Revised")
    await update.message.reply_text(option_b_text, parse_mode="Markdown")

    keyboard = [
        [
            InlineKeyboardButton("Pick A", callback_data="approve_a"),
            InlineKeyboardButton("Pick B", callback_data="approve_b"),
        ],
        [
            InlineKeyboardButton("Revise A", callback_data="approve_revise_a"),
            InlineKeyboardButton("Revise B", callback_data="approve_revise_b"),
        ],
        [
            InlineKeyboardButton("Regenerate", callback_data="approve_regenerate"),
            InlineKeyboardButton("Cancel", callback_data="approve_cancel"),
        ],
    ]
    await update.message.reply_text(
        "How about these?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return POST_APPROVE


# ============================================================
# Calendar flow
# ============================================================

async def calendar_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start the weekly calendar flow."""
    message = update.callback_query.message if update.callback_query else update.message
    if update.callback_query:
        await update.callback_query.answer()

    await message.reply_text(
        "Let's plan your content calendar.\n\n"
        "What are your themes this week? List them separated by commas.\n\n"
        "Example: uncertainty, career transitions, self-compassion"
    )
    return CAL_THEMES


async def calendar_themes(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle theme input for calendar."""
    themes_text = update.message.text
    context.user_data["cal_themes"] = [t.strip() for t in themes_text.split(",")]

    keyboard = [
        [InlineKeyboardButton("Instagram", callback_data="calplatform_instagram")],
        [InlineKeyboardButton("Rednote", callback_data="calplatform_rednote")],
        [InlineKeyboardButton("Both", callback_data="calplatform_both")],
    ]
    await update.message.reply_text(
        "Generate calendar for which platform?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return CAL_PLATFORM


async def calendar_platform(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle platform selection for calendar."""
    query = update.callback_query
    await query.answer()

    platform = query.data.replace("calplatform_", "")
    context.user_data["cal_platform"] = platform

    keyboard = [
        [InlineKeyboardButton("3 posts", callback_data="calfreq_3")],
        [InlineKeyboardButton("5 posts", callback_data="calfreq_5")],
        [InlineKeyboardButton("7 posts", callback_data="calfreq_7")],
    ]
    await query.message.reply_text(
        "How many posts this week?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return CAL_FREQUENCY


async def calendar_frequency(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle frequency selection and generate calendar."""
    query = update.callback_query
    await query.answer()

    num_posts = int(query.data.replace("calfreq_", ""))
    context.user_data["cal_num_posts"] = num_posts

    await query.message.reply_text(
        f"Generating {num_posts}-post calendar... This may take a minute."
    )

    engine = _get_content_engine()
    platform = context.user_data["cal_platform"]
    themes = context.user_data["cal_themes"]

    platforms = [platform] if platform != "both" else ["instagram", "rednote"]
    calendar_text = engine.generate_calendar(
        platforms=platforms,
        themes=themes,
        num_posts=num_posts,
    )
    context.user_data["calendar_text"] = calendar_text
    context.user_data["cal_current_day"] = 0

    # Send the full calendar (split into chunks if needed)
    await _send_calendar_chunks(query.message, calendar_text)

    keyboard = [
        [InlineKeyboardButton("Approve All", callback_data="calnav_approve_all")],
        [InlineKeyboardButton("Done", callback_data="calnav_done")],
    ]
    await query.message.reply_text(
        "Here's your weekly calendar. You can approve it or make changes.",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )
    return CAL_NAVIGATE


async def calendar_navigate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle calendar navigation and approval."""
    query = update.callback_query
    await query.answer()

    action = query.data.replace("calnav_", "")

    if action in ("approve_all", "done"):
        context.user_data.clear()
        await query.message.reply_text(
            "Calendar saved! You can copy each post when you're ready to publish.\n\n"
            "Send /start to create more content."
        )
        return ConversationHandler.END


async def _send_calendar_chunks(message, text: str, chunk_size: int = 4000):
    """Split long calendar text into Telegram-friendly chunks."""
    if len(text) <= chunk_size:
        await message.reply_text(text)
        return

    # Split on day boundaries
    parts = text.split("====== DAY")
    for i, part in enumerate(parts):
        if not part.strip():
            continue
        chunk = f"====== DAY{part}" if i > 0 else part
        if len(chunk) > chunk_size:
            # Further split if needed
            for j in range(0, len(chunk), chunk_size):
                await message.reply_text(chunk[j : j + chunk_size])
        else:
            await message.reply_text(chunk)


# ============================================================
# Brand voice update
# ============================================================

async def update_voice_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Start the brand voice update flow."""
    await update.message.reply_text(
        "Upload your new brand voice document (.docx format).\n\n"
        "I'll back up the current one and start using the new guide immediately."
    )
    return VOICE_UPLOAD


async def update_voice_upload(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle brand voice document upload."""
    doc = update.message.document
    if not doc.file_name.endswith(".docx"):
        await update.message.reply_text(
            "Please upload a .docx file. Other formats aren't supported yet."
        )
        return VOICE_UPLOAD

    try:
        file = await doc.get_file()
        file_bytes = await file.download_as_bytearray()

        # Convert .docx to text
        docx_doc = DocxDocument(io.BytesIO(bytes(file_bytes)))
        lines = []
        for paragraph in docx_doc.paragraphs:
            text = paragraph.text.strip()
            lines.append(text if text else "")
        for table in docx_doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
            lines.append("")

        new_content = "\n".join(lines)

        # Backup current brand voice
        if BRAND_VOICE_FILE.exists():
            shutil.copy2(BRAND_VOICE_FILE, BRAND_VOICE_BACKUP)

        # Write new brand voice
        BRAND_VOICE_FILE.write_text(new_content, encoding="utf-8")

        await update.message.reply_text(
            "Brand voice updated!\n\n"
            f"New guide: {len(new_content)} characters, {len(new_content.split())} words.\n"
            "Your next content will use the new guide automatically."
        )

    except Exception as e:
        logger.error(f"Brand voice update failed: {e}")
        await update.message.reply_text(
            f"Something went wrong processing the document. Error: {e}\n"
            "Please try again or contact support."
        )

    return ConversationHandler.END


# ============================================================
# Free-form message handling
# ============================================================

async def free_form_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle natural language messages outside of structured flows."""
    text = update.message.text.lower()

    # Detect platform intent
    platform = "instagram"  # default
    if "rednote" in text or "xiaohongshu" in text or "小红书" in text:
        platform = "rednote"
    elif "both" in text:
        platform = "both"

    # Use the full message as the topic
    context.user_data["platform"] = platform
    context.user_data["topic"] = update.message.text
    context.user_data["image_choice"] = "skip"

    await update.message.reply_text("On it! Generating 2 options for you...")

    engine = _get_content_engine()
    result = engine.generate_post(platform=platform, topic=update.message.text)
    context.user_data["result"] = result
    context.user_data["conversation_history"] = [
        {"role": "user", "content": f"Topic: {update.message.text}"},
        {"role": "assistant", "content": result["raw"]},
    ]

    # Send options
    option_a_text = _format_option_preview(result["option_a"], "Option A (Reflective)")
    await update.message.reply_text(option_a_text, parse_mode="Markdown")

    option_b_text = _format_option_preview(result["option_b"], "Option B (Direct)")
    await update.message.reply_text(option_b_text, parse_mode="Markdown")

    keyboard = [
        [
            InlineKeyboardButton("Pick A", callback_data="approve_a"),
            InlineKeyboardButton("Pick B", callback_data="approve_b"),
        ],
        [
            InlineKeyboardButton("Revise A", callback_data="approve_revise_a"),
            InlineKeyboardButton("Revise B", callback_data="approve_revise_b"),
        ],
        [
            InlineKeyboardButton("Regenerate", callback_data="approve_regenerate"),
        ],
    ]
    await update.message.reply_text(
        "Which one do you prefer?",
        reply_markup=InlineKeyboardMarkup(keyboard),
    )


# ============================================================
# Helpers
# ============================================================

def _format_option_preview(option: dict, label: str) -> str:
    """Format an option for Telegram display."""
    lines = [f"*{label}*"]
    if option.get("content_type"):
        lines.append(f"Type: {option['content_type']}")
    lines.append("")

    # Use raw if structured parsing didn't work well
    content = option.get("caption") or option.get("raw", "")
    if content:
        lines.append(content)

    if option.get("hashtags"):
        lines.append("")
        lines.append(option["hashtags"])

    if option.get("visual"):
        lines.append("")
        lines.append(f"Visual: {option['visual']}")

    return "\n".join(lines)
